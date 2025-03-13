"""
Advanced Asynchronous Web Scraper
--------------------------------
A modular, high-performance web scraper with fault tolerance,
resource management, and comprehensive content analysis.

Usage:
    python scraper.py [--config CONFIG_FILE]
"""

import asyncio
import hashlib
import json
import logging
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from threading import Lock
from typing import Any, Dict, List, Optional, Set, Tuple, Union
from urllib import robotparser
from urllib.parse import urljoin, urlparse

import aiohttp
import psutil
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# Optional imports with proper error handling
try:
    from playwright.async_api import async_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False

try:
    from transformers import AutoModelForCausalLM, AutoTokenizer
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False


@dataclass
class ScraperConfig:
    """Configuration for the web scraper with comprehensive documentation.
    
    This class centralizes all parameters that control the behavior of the scraper,
    ensuring type safety and providing default values where appropriate.
    
    Attributes:
        root_url: Starting URL for the scraping process
        output_dir: Directory where scraped content and reports will be stored
        max_retries: Maximum number of retry attempts for failed requests
        delay: Base delay between requests in seconds (will be increased for retries)
        max_workers: Maximum number of concurrent scraping tasks
        timeout: Request timeout in seconds
        max_depth: Maximum link traversal depth from root URL
        user_agent: User agent string to identify the scraper to web servers
        excluded_patterns: URL patterns to exclude from scraping
        memory_limit_mb: Maximum memory usage allowed in MB
        enable_javascript: Whether to render JavaScript content using browser automation
        proxies: List of proxy URLs for request rotation
        max_file_size_mb: Maximum file size to process in MB
    """
    
    root_url: str
    output_dir: Path
    max_retries: int = 3
    delay: float = 1.0
    max_workers: int = 5
    timeout: int = 30
    max_depth: int = 5
    user_agent: str = "Advanced Web Scraper 2.0"
    excluded_patterns: List[str] = field(default_factory=lambda: [
        ".pdf", ".zip", "#", "mailto:", "javascript:"
    ])
    memory_limit_mb: int = 1000
    enable_javascript: bool = False
    proxies: List[str] = field(default_factory=list)
    max_file_size_mb: int = 50
    content_type_filters: List[str] = field(default_factory=lambda: [
        "text/html", "application/xhtml+xml"
    ])
    extract_metadata: bool = True
    follow_redirects: bool = True
    verify_ssl: bool = True
    language: str = "en"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert configuration to dictionary format.
        
        Returns:
            Dictionary representation of configuration
        """
        return asdict(self)
    
    @classmethod
    def from_dict(cls, config_dict: Dict[str, Any]) -> 'ScraperConfig':
        """Create configuration from dictionary.
        
        Args:
            config_dict: Dictionary containing configuration parameters
            
        Returns:
            ScraperConfig instance with parameters from dictionary
        """
        # Convert string path to Path object
        if "output_dir" in config_dict and isinstance(config_dict["output_dir"], str):
            config_dict["output_dir"] = Path(config_dict["output_dir"])
        
        return cls(**config_dict)
    
    @classmethod
    def from_json_file(cls, file_path: Union[str, Path]) -> 'ScraperConfig':
        """Load configuration from JSON file.
        
        Args:
            file_path: Path to JSON configuration file
            
        Returns:
            ScraperConfig instance with parameters from file
            
        Raises:
            FileNotFoundError: If the specified file doesn't exist
            json.JSONDecodeError: If the file contains invalid JSON
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            config_dict = json.load(f)
        
        return cls.from_dict(config_dict)
    
    def save_to_json(self, file_path: Union[str, Path]) -> None:
        """Save configuration to JSON file.
        
        Args:
            file_path: Path where to save the configuration
        """
        config_dict = self.to_dict()
        # Convert Path to string for JSON serialization
        config_dict["output_dir"] = str(config_dict["output_dir"])
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(config_dict, f, indent=2, ensure_ascii=False)


class StorageBackend:
    """Abstract base class for storage implementations.
    
    This class defines the interface that all storage backends must implement
    to provide consistent storage capabilities to the scraper.
    """
    
    def save_content(self, url: str, content: str) -> None:
        """Save scraped content for a given URL.
        
        Args:
            url: The URL the content was scraped from
            content: The scraped content to save
            
        Raises:
            NotImplementedError: If the subclass does not implement this method
        """
        raise NotImplementedError
    
    def save_urls(self, urls: Set[str]) -> None:
        """Save a set of URLs that have been processed.
        
        Args:
            urls: Set of URLs to save
            
        Raises:
            NotImplementedError: If the subclass does not implement this method
        """
        raise NotImplementedError
    
    def load_urls(self) -> Set[str]:
        """Load previously processed URLs.
        
        Returns:
            Set of previously processed URLs
            
        Raises:
            NotImplementedError: If the subclass does not implement this method
        """
        raise NotImplementedError
    
    def should_process_url(self, url: str, content: Optional[str] = None) -> bool:
        """Check if URL should be processed based on previous scraping history.
        
        Args:
            url: URL to check
            content: Optional content to compare with previous version
            
        Returns:
            Boolean indicating whether the URL should be processed
            
        Raises:
            NotImplementedError: If the subclass does not implement this method
        """
        raise NotImplementedError


class FileStorage(StorageBackend):
    """File-based storage backend with advanced organization and metadata.
    
    This class implements the StorageBackend interface using the file system
    for persistent storage of scraped content, with domain-specific organization
    and content deduplication.
    
    Attributes:
        output_dir: Base directory for all storage
        content_dir: Subdirectory for storing content files
        metadata_dir: Subdirectory for storing metadata
        logs_dir: Subdirectory for storing logs
    """
    
    def __init__(self, output_dir: Path):
        """Initialize file storage with appropriate directory structure.
        
        Args:
            output_dir: Base directory for all storage
        """
        self.output_dir = output_dir
        self.output_dir.mkdir(exist_ok=True)
        
        # Create subdirectories for different types of content
        self.content_dir = self.output_dir / "content"
        self.metadata_dir = self.output_dir / "metadata"
        self.logs_dir = self.output_dir / "logs"
        
        for directory in [self.content_dir, self.metadata_dir, self.logs_dir]:
            directory.mkdir(exist_ok=True)
            
        self._url_lock = Lock()
        self._content_lock = Lock()
        self._metadata = self._load_metadata()
        
    def _load_metadata(self) -> Dict[str, Dict[str, Any]]:
        """Load metadata including URL hashes and additional information.
        
        Returns:
            Dictionary mapping URLs to their metadata
        """
        try:
            with (self.metadata_dir / "metadata.json").open("r", encoding="utf-8") as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            return {}
            
    def _save_metadata(self) -> None:
        """Save metadata with proper JSON formatting."""
        with (self.metadata_dir / "metadata.json").open("w", encoding="utf-8") as f:
            json.dump(self._metadata, f, ensure_ascii=False, indent=2, sort_keys=True)
            
    def _get_domain_path(self, url: str) -> Path:
        """Get domain-specific directory path.
        
        Args:
            url: URL to extract domain from
            
        Returns:
            Path object for domain-specific directory
        """
        domain = urlparse(url).netloc
        domain_dir = self.content_dir / domain
        domain_dir.mkdir(exist_ok=True)
        return domain_dir
        
    def save_content(self, url: str, content: str) -> None:
        """Save content with improved organization and validation.
        
        Args:
            url: The URL the content was scraped from
            content: The scraped content to save
        """
        content_hash = hashlib.md5(content.encode()).hexdigest()
        
        # Skip if content hasn't changed
        if url in self._metadata and self._metadata[url].get("content_hash") == content_hash:
            return
            
        with self._content_lock:
            # Create domain-specific directory
            domain_dir = self._get_domain_path(url)
            
            # Create dated filename
            date_str = datetime.now().strftime("%Y%m%d")
            filename = f"{date_str}_{hashlib.md5(url.encode()).hexdigest()}.json"
            file_path = domain_dir / filename
            
            # Prepare content with metadata
            data = {
                "url": url,
                "content": content,
                "timestamp": datetime.now().isoformat(),
                "content_hash": content_hash,
                "metadata": {
                    "word_count": len(content.split()),
                    "char_count": len(content),
                    "date_processed": date_str
                }
            }
            
            # Save content with proper JSON formatting
            with file_path.open("w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2, sort_keys=True)
                
            # Update metadata
            self._metadata[url] = {
                "content_hash": content_hash,
                "file_path": str(file_path.relative_to(self.output_dir)),
                "last_updated": datetime.now().isoformat()
            }
            
            # Save updated metadata
            self._save_metadata()
            
    def save_urls(self, urls: Set[str]) -> None:
        """Save processed URLs with metadata.
        
        Args:
            urls: Set of URLs to save
        """
        with self._url_lock:
            # Update processed URLs in metadata
            for url in urls:
                if url not in self._metadata:
                    self._metadata[url] = {
                        "first_seen": datetime.now().isoformat(),
                        "times_processed": 1
                    }
                else:
                    self._metadata[url]["times_processed"] = \
                        self._metadata[url].get("times_processed", 0) + 1
                    
            self._save_metadata()
            
    def load_urls(self) -> Set[str]:
        """Load previously processed URLs from metadata.
        
        Returns:
            Set of previously processed URLs
        """
        return set(self._metadata.keys())
        
    def should_process_url(self, url: str, content: Optional[str] = None) -> bool:
        """Check if URL should be processed based on content hash.
        
        Args:
            url: URL to check
            content: Optional content to compare with previous version
            
        Returns:
            Boolean indicating whether the URL should be processed
        """
        if content is None:
            return url not in self._metadata
            
        content_hash = hashlib.md5(content.encode()).hexdigest()
        return url not in self._metadata or \
               self._metadata[url].get("content_hash") != content_hash


class ContentProcessor:
    """Handles text processing and content management with optional AI enhancement.
    
    This class manages text content, buffers it until reaching size limits,
    and optionally processes it using transformer models for enhancement.
    
    Attributes:
        char_limit: Maximum characters per content buffer
        current_text: Current buffer of accumulated text
        file_counter: Counter for generated files
        model: Optional transformer model for text processing
        tokenizer: Optional tokenizer for transformer model
        model_name: Name of transformer model to use
    """
    
    def __init__(self, 
                 char_limit: int = 1000000, 
                 enable_ai_processing: bool = False,
                 model_name: str = "ibm-granite/granite-3.1-8b-instruct"):
        """Initialize content processor.
        
        Args:
            char_limit: Maximum characters per content buffer
            enable_ai_processing: Whether to enable AI-based text processing
            model_name: Name of transformer model to use if AI processing is enabled
        """
        self.char_limit = char_limit
        self.current_text = ""
        self.file_counter = 1
        self.model = None
        self.tokenizer = None
        self.model_name = model_name
        
        # Only attempt to load models if AI processing is enabled and transformers are available
        if enable_ai_processing and TRANSFORMERS_AVAILABLE:
            try:
                self.tokenizer = AutoTokenizer.from_pretrained(self.model_name, local_files_only=True)
                self.model = AutoModelForCausalLM.from_pretrained(self.model_name, local_files_only=True)
                logging.info(f"Model successfully loaded from local cache: {model_name}")
            except Exception as e:
                logging.warning(f"Cannot load model from local cache: {e}")
                try:
                    self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
                    self.model = AutoModelForCausalLM.from_pretrained(self.model_name)
                    logging.info(f"Model successfully downloaded and loaded: {model_name}")
                except Exception as e:
                    logging.error(f"Failed to load model: {e}")
                    logging.warning("Continuing without model processing")
        elif enable_ai_processing and not TRANSFORMERS_AVAILABLE:
            logging.warning("AI processing requested but transformers library not available")
        
    def process_text(self, text: str) -> str:
        """Process text using transformer model if available, otherwise return original text.
        
        Args:
            text: Input text to process
            
        Returns:
            Processed or original text
        """
        if self.model is None or self.tokenizer is None:
            return text
            
        try:
            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
            outputs = self.model.generate(**inputs, max_length=150)
            processed_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            return processed_text
        except Exception as e:
            logging.error(f"Error processing text with model: {e}")
            return text
        
    def should_create_new_file(self, new_text: str) -> bool:
        """Check if adding new text would exceed character limit.
        
        Args:
            new_text: Text to be added
            
        Returns:
            Boolean indicating whether a new file should be created
        """
        return len(self.current_text) + len(new_text) > self.char_limit
        
    def add_text(self, text: str, title: str) -> Optional[str]:
        """Add text to current buffer, return buffer if limit reached.
        
        Args:
            text: Text content to add
            title: Title for the content (typically URL)
            
        Returns:
            Current buffer content if limit reached, otherwise None
        """
        formatted_text = f"{title}\n{text}\n\n"
        
        if self.should_create_new_file(formatted_text):
            old_text = self.current_text
            self.current_text = formatted_text
            self.file_counter += 1
            return old_text
            
        self.current_text += formatted_text
        return None
        
    def get_current_buffer(self) -> str:
        """Get current text buffer.
        
        Returns:
            Current accumulated text
        """
        return self.current_text
        
    def clear_buffer(self) -> None:
        """Clear the current text buffer."""
        self.current_text = ""


class RateLimiter:
    """Rate limiter implementation using token bucket algorithm.
    
    This class implements a token bucket algorithm for precise rate limiting,
    ensuring polite crawling by controlling request frequency.
    
    Attributes:
        calls: Maximum number of calls allowed per period
        period: Time period in seconds
        tokens: Current number of available tokens
        last_update: Timestamp of last token update
    """
    
    def __init__(self, calls: int, period: float):
        """Initialize rate limiter.
        
        Args:
            calls: Maximum number of calls allowed per period
            period: Time period in seconds
        """
        self.calls = calls
        self.period = period
        self.tokens = calls
        self.last_update = time.time()
        self._lock = asyncio.Lock()

    async def acquire(self) -> None:
        """Acquire permission to make a request.
        
        Implements token bucket algorithm for more precise rate limiting.
        """
        async with self._lock:
            now = time.time()
            time_passed = now - self.last_update
            self.tokens = min(self.calls, self.tokens + time_passed * (self.calls / self.period))
            self.last_update = now

            if self.tokens < 1:
                sleep_time = (1 - self.tokens) * (self.period / self.calls)
                await asyncio.sleep(sleep_time)
                self.tokens = 0
            else:
                self.tokens -= 1


class CircuitBreaker:
    """Circuit breaker pattern implementation for handling failing domains.
    
    This class implements the circuit breaker design pattern to prevent
    repeated failures when accessing problematic domains.
    
    Attributes:
        failure_threshold: Number of failures before opening circuit
        reset_timeout: Time in seconds before attempting reset
        failures: Dictionary tracking failures per domain
        last_failure_time: Dictionary tracking last failure time per domain
    """
    
    def __init__(self, failure_threshold: int = 5, reset_timeout: float = 300):
        """Initialize circuit breaker.
        
        Args:
            failure_threshold: Number of failures before opening circuit
            reset_timeout: Time in seconds before attempting reset
        """
        self.failure_threshold = failure_threshold
        self.reset_timeout = reset_timeout
        self.failures: Dict[str, int] = {}
        self.last_failure_time: Dict[str, float] = {}
        self._lock = asyncio.Lock()

    async def record_failure(self, domain: str) -> None:
        """Record a failure for a domain.
        
        Args:
            domain: Domain that experienced a failure
        """
        async with self._lock:
            self.failures[domain] = self.failures.get(domain, 0) + 1
            self.last_failure_time[domain] = time.time()

    async def record_success(self, domain: str) -> None:
        """Record a success for a domain.
        
        Args:
            domain: Domain that experienced a success
        """
        async with self._lock:
            if domain in self.failures:
                del self.failures[domain]
                del self.last_failure_time[domain]

    async def is_open(self, domain: str) -> bool:
        """Check if circuit is open for domain.
        
        Args:
            domain: Domain to check
            
        Returns:
            Boolean indicating whether the circuit is open
        """
        async with self._lock:
            if domain not in self.failures:
                return False

            if self.failures[domain] >= self.failure_threshold:
                # Check if enough time has passed to try again
                if time.time() - self.last_failure_time[domain] > self.reset_timeout:
                    del self.failures[domain]
                    del self.last_failure_time[domain]
                    return False
                return True
            return False


class MemoryMonitor:
    """Monitor and manage memory usage during scraping.
    
    This class monitors memory usage to prevent the scraper from
    consuming excessive system resources.
    
    Attributes:
        limit_bytes: Maximum memory usage allowed in bytes
        process: Current process for memory monitoring
    """
    
    def __init__(self, limit_mb: int):
        """Initialize memory monitor.
        
        Args:
            limit_mb: Maximum memory usage allowed in MB
        """
        self.limit_bytes = limit_mb * 1024 * 1024
        self.process = psutil.Process()
        
    def check_memory(self) -> Tuple[bool, float]:
        """Check current memory usage.
        
        Returns:
            Tuple of (is_within_limit, current_usage_mb)
        """
        memory_info = self.process.memory_info()
        current_usage = memory_info.rss
        return current_usage < self.limit_bytes, current_usage / (1024 * 1024)


class ProxyManager:
    """Manage and rotate proxies for scraping.
    
    This class provides proxy rotation functionality to distribute
    requests across multiple proxies for better anonymity and
    load balancing.
    
    Attributes:
        proxies: List of proxy URLs
        current_index: Current index in proxy rotation
    """
    
    def __init__(self, proxies: List[str]):
        """Initialize proxy manager.
        
        Args:
            proxies: List of proxy URLs
        """
        self.proxies = proxies
        self.current_index = 0
        self._lock = asyncio.Lock()
        
    async def get_next_proxy(self) -> Optional[str]:
        """Get next proxy in rotation.
        
        Returns:
            Proxy URL or None if no proxies configured
        """
        if not self.proxies:
            return None
            
        async with self._lock:
            proxy = self.proxies[self.current_index]
            self.current_index = (self.current_index + 1) % len(self.proxies)
            return proxy


class PageFetcher:
    """Component responsible for fetching web pages with JavaScript support.
    
    This class handles the actual HTTP requests and optional JavaScript
    rendering using Playwright.
    
    Attributes:
        config: Scraper configuration
        playwright: Playwright instance for browser automation
        browser: Browser instance for JavaScript rendering
    """
    
    def __init__(self, config: ScraperConfig):
        """Initialize page fetcher.
        
        Args:
            config: Scraper configuration
        """
        self.config = config
        self.playwright = None
        self.browser = None
        
    async def setup(self):
        """Initialize browser if JavaScript is enabled."""
        if self.config.enable_javascript and PLAYWRIGHT_AVAILABLE:
            try:
                self.playwright = await async_playwright().start()
                self.browser = await self.playwright.chromium.launch()
                logging.info("Playwright browser initialized for JavaScript rendering")
            except Exception as e:
                logging.error(f"Failed to initialize Playwright: {e}")
                logging.warning("Continuing without JavaScript support")
        elif self.config.enable_javascript and not PLAYWRIGHT_AVAILABLE:
            logging.warning("JavaScript rendering requested but Playwright not available")
            
    async def cleanup(self):
        """Cleanup browser resources."""
        if self.browser:
            await self.browser.close()
        if self.playwright:
            await self.playwright.stop()
            
    async def fetch_page(
        self, url: str, session: aiohttp.ClientSession, headers: Dict[str, str]
    ) -> Optional[str]:
        """Fetch page content with optional JavaScript rendering.
        
        Args:
            url: URL to fetch
            session: aiohttp session for HTTP requests
            headers: HTTP headers to include in request
            
        Returns:
            Page content as string or None if fetch failed
        """
        if self.config.enable_javascript and self.browser:
            try:
                page = await self.browser.new_page()
                await page.set_extra_http_headers(headers)
                response = await page.goto(
                    url, 
                    wait_until='networkidle',
                    timeout=self.config.timeout * 1000  # Convert to milliseconds
                )
                
                if not response or response.status != 200:
                    logging.warning(f"Failed to load page with JavaScript: {url}, status: {response.status if response else 'No response'}")
                    await page.close()
                    return None
                
                content = await page.content()
                await page.close()
                return content
            except Exception as e:
                logging.error(f"Error rendering JavaScript for {url}: {e}")
                return None
        else:
            try:
                async with session.get(
                    url, 
                    headers=headers, 
                    timeout=self.config.timeout,
                    ssl=self.config.verify_ssl,
                    allow_redirects=self.config.follow_redirects
                ) as response:
                    if response.status != 200:
                        logging.warning(f"Failed to fetch {url}, status: {response.status}")
                        return None
                    
                    # Check content type if specified in config
                    if self.config.content_type_filters:
                        content_type = response.headers.get('Content-Type', '')
                        if not any(filter_type in content_type for filter_type in self.config.content_type_filters):
                            logging.warning(f"Skipping {url}, content type not allowed: {content_type}")
                            return None
                    
                    return await response.text()
            except Exception as e:
                logging.error(f"Error fetching {url}: {e}")
                return None


class URLManager:
    """Component for managing URL validation and tracking.
    
    This class handles URL validation, robots.txt compliance, and
    sitemap parsing to ensure proper URL traversal.
    
    Attributes:
        config: Scraper configuration
        visited_urls: Set of already visited URLs
        robots_parser: Parser for robots.txt rules
        sitemap_urls: URLs extracted from sitemap.xml
    """
    
    def __init__(self, config: ScraperConfig):
        """Initialize URL manager.
        
        Args:
            config: Scraper configuration
        """
        self.config = config
        self.visited_urls: Set[str] = set()
        self.robots_parser = robotparser.RobotFileParser()
        self.sitemap_urls: Set[str] = set()
        self._setup_robots()
        self._setup_sitemap()
        
    def _setup_robots(self) -> None:
        """Setup robots.txt parser."""
        try:
            root_parts = urlparse(self.config.root_url)
            if root_parts.scheme and root_parts.netloc:
                robots_url = f"{root_parts.scheme}://{root_parts.netloc}/robots.txt"
                self.robots_parser.set_url(robots_url)
                self.robots_parser.read()
                logging.info(f"Robots.txt loaded from {robots_url}")
        except Exception as e:
            logging.warning(f"Failed to load robots.txt: {e}")
            
    def _setup_sitemap(self) -> None:
        """Load sitemap.xml."""
        try:
            root_parts = urlparse(self.config.root_url)
            if root_parts.scheme and root_parts.netloc:
                sitemap_url = f"{root_parts.scheme}://{root_parts.netloc}/sitemap.xml"
                response = requests.get(
                    sitemap_url, 
                    timeout=self.config.timeout,
                    verify=self.config.verify_ssl
                )
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, 'xml')
                    self.sitemap_urls = {loc.text for loc in soup.find_all('loc')}
                    logging.info(f"Loaded {len(self.sitemap_urls)} URLs from sitemap.xml")
        except Exception as e:
            logging.warning(f"Failed to load sitemap: {e}")
            
    def is_valid_url(self, url: str, depth: int) -> bool:
        """Check if URL should be processed.
        
        Args:
            url: URL to validate
            depth: Current depth from root URL
            
        Returns:
            Boolean indicating whether the URL is valid for processing
        """
        if depth > self.config.max_depth:
            return False
            
        try:
            result = urlparse(url)
            if not all([result.scheme, result.netloc]) or \
                not url.startswith(self.config.root_url):
                return False
                
            if not self.robots_parser.can_fetch(self.config.user_agent, url):
                logging.info(f"Skipping {url} - disallowed by robots.txt")
                return False
                
            if self.config.excluded_patterns and \
                any(pattern in url for pattern in self.config.excluded_patterns):
                return False
                
            return True
        except ValueError:
            return False
            
    def mark_visited(self, url: str) -> None:
        """Mark URL as visited.
        
        Args:
            url: URL to mark as visited
        """
        self.visited_urls.add(url)
        
    def is_visited(self, url: str) -> bool:
        """Check if URL has been visited.
        
        Args:
            url: URL to check
            
        Returns:
            Boolean indicating whether the URL has been visited
        """
        return url in self.visited_urls
        
    def get_sitemap_links(self, url: str, depth: int) -> List[Tuple[str, str]]:
        """Get relevant sitemap links for URL.
        
        Args:
            url: URL to match against sitemap
            depth: Current depth from root URL
            
        Returns:
            List of tuples containing (url, text) for matching sitemap links
        """
        return [
            (sitemap_url, "")
            for sitemap_url in self.sitemap_urls
            if sitemap_url.startswith(url) and self.is_valid_url(sitemap_url, depth + 1)
        ]


class RealTimeContentAnalyzer:
    """Analyzes content in real-time during scraping.
    
    This class generates analysis reports and content summaries
    during the scraping process.
    
    Attributes:
        storage_dir: Directory for storing analysis results
        document: Word document for content analysis
        content_map: Mapping of content categories
        all_contents: List of all processed contents
        common_prefix: Common prefix found in contents
        common_suffix: Common suffix found in contents
        url_structure: Dictionary counting URL patterns
        content_stats: Statistics about content
    """
    
    def __init__(self, storage_dir: Path):
        """Initialize content analyzer.
        
        Args:
            storage_dir: Directory for storing analysis results
        """
        self.storage_dir = storage_dir
        self.document = Document()
        self.content_map = {}
        self.all_contents = []
        self.common_prefix = None
        self.common_suffix = None
        self.url_structure = {}
        self.content_stats = {
            'total_words': 0,
            'avg_length': 0,
            'processed_count': 0
        }
        # Patterns for cleaning
        self.navigation_patterns = [
            "Search documentation...",
            "Search... ⌘K",
            "Skip to content",
            "Showcase Docs Blog Templates Enterprise",
        ]
        self.menu_start = "Using Latest Version"
        self.menu_end = "TypeScript ESLint CLI"
        
        # Initialize document structure
        self._setup_document()

    def _setup_document(self):
        """Initialize document structure."""
        self.document.styles["Normal"].font.name = "Calibri"
        self.document.styles["Normal"].font.size = Pt(11)
        
        # Title page
        title = self.document.add_heading("Real-Time Content Analysis", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = self.document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(datetime.now().strftime("%Y-%m-%d")).bold = True
        
        self.document.add_page_break()
        
    def process_content(self, url: str, content: str) -> None:
        """Process content in real-time during scraping.
        
        Args:
            url: URL of the content
            content: Text content to analyze
        """
        try:
            # Clean content
            cleaned_content = self.clean_content(content)
            if not cleaned_content.strip():
                return
                
            # Update statistics
            words = cleaned_content.split()
            self.content_stats['total_words'] += len(words)
            self.content_stats['processed_count'] += 1
            self.content_stats['avg_length'] = (
                self.content_stats['total_words'] / self.content_stats['processed_count']
            )
            
            # Analyze URL structure
            url_parts = urlparse(url).path.split("/")
            if url_parts:
                path_key = url_parts[1] if len(url_parts) > 1 else url_parts[0]
                self.url_structure[path_key] = self.url_structure.get(path_key, 0) + 1
                
            # Add to content map
            path_key = url_parts[1] if len(url_parts) > 1 else "root"
            if path_key not in self.content_map:
                self.content_map[path_key] = []
                
            self.content_map[path_key].append({
                "url": url,
                "content": cleaned_content,
                "timestamp": datetime.now().isoformat()
            })
            
            # Add content to analysis list
            self.all_contents.append(content)
            
            # Update document in real-time
            self._update_document()
            
            # Save progress periodically
            if self.content_stats['processed_count'] % 10 == 0:
                self._find_common_patterns()
                self.save_document()
                
        except Exception as e:
            logging.error(f"Error processing content for {url}: {e}")

    def _find_common_patterns(self):
        """Find common patterns in content for cleaning."""
        if len(self.all_contents) < 2:
            return

        # Find common prefix
        min_length = min(len(content) for content in self.all_contents)
        
        # Find maximum common prefix
        prefix_length = 0
        for i in range(min_length):
            if all(content[i] == self.all_contents[0][i] for content in self.all_contents):
                prefix_length = i + 1
            else:
                break
                
        # Find maximum common suffix
        suffix_length = 0
        for i in range(min_length):
            pos = -1 - i
            if all(content[pos] == self.all_contents[0][pos] for content in self.all_contents):
                suffix_length = i + 1
            else:
                break
                
        # Set common patterns if they're substantial
        if prefix_length > 100:
            self.common_prefix = self.all_contents[0][:prefix_length]
            logging.info(f"Found common prefix of length {prefix_length}")
            
        if suffix_length > 100:
            self.common_suffix = self.all_contents[0][-suffix_length:]
            logging.info(f"Found common suffix of length {suffix_length}")

    def clean_content(self, text: str) -> str:
        """Remove navigation elements and repeating parts from text.
        
        Args:
            text: Text to clean
            
        Returns:
            Cleaned text
        """
        # First remove common parts
        if self.common_prefix and text.startswith(self.common_prefix):
            text = text[len(self.common_prefix):]
        if self.common_suffix and text.endswith(self.common_suffix):
            text = text[:-len(self.common_suffix)]

        # Process text line by line
        lines = text.split("\n")
        cleaned_lines = []
        inside_menu = False

        for line in lines:
            line = line.strip()

            # Skip short lines
            if len(line) < 10:
                continue

            # Skip navigation patterns
            if any(pattern in line for pattern in self.navigation_patterns):
                continue

            # Detect menu sections
            if self.menu_start in line:
                inside_menu = True
                continue

            if self.menu_end in line:
                inside_menu = False
                continue

            # Skip lines with many links
            if line.count("|") > 3 or line.count(">") > 3:
                continue

            if not inside_menu:
                cleaned_lines.append(line)

        # Join lines and remove duplicate empty lines
        cleaned_text = "\n".join(cleaned_lines)
        cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
        return cleaned_text.strip()

    def _update_document(self):
        """Update the analysis document with current data."""
        # Clear existing document
        self.document = Document()
        self._setup_document()
        
        # Add statistics section
        self.document.add_heading("Content Statistics", 1)
        stats_table = self.document.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        
        # Add statistics rows
        stats = [
            ("Total Content Items", str(self.content_stats['processed_count'])),
            ("Total Words", str(self.content_stats['total_words'])),
            ("Average Words per Item", f"{self.content_stats['avg_length']:.1f}"),
            ("Content Categories", str(len(self.content_map)))
        ]
        
        for i, (key, value) in enumerate(stats):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = value
            
        self.document.add_paragraph()
        
        # Add URL structure section
        self.document.add_heading("URL Structure Analysis", 1)
        if self.url_structure:
            url_table = self.document.add_table(rows=len(self.url_structure) + 1, cols=2)
            url_table.style = 'Table Grid'
            
            # Header row
            url_table.cell(0, 0).text = "Path Component"
            url_table.cell(0, 1).text = "Count"
            
            # Data rows
            for i, (path, count) in enumerate(sorted(self.url_structure.items(), key=lambda x: x[1], reverse=True)):
                url_table.cell(i + 1, 0).text = path
                url_table.cell(i + 1, 1).text = str(count)
        else:
            self.document.add_paragraph("No URL structure data available yet.")
            
        self.document.add_page_break()
        
        # Add content categories section
        self.document.add_heading("Content Categories", 1)
        for category, items in self.content_map.items():
            self.document.add_heading(category.capitalize(), 2)
            self.document.add_paragraph(f"Contains {len(items)} items")
            
            # Add sample content if available
            if items:
                sample = items[0]
                self.document.add_heading("Sample URL", 3)
                self.document.add_paragraph(sample["url"])
                
                self.document.add_heading("Sample Content", 3)
                # Add at most first 500 characters to keep document concise
                content_text = sample["content"][:500] + ("..." if len(sample["content"]) > 500 else "")
                self.document.add_paragraph(content_text)
                
            self.document.add_page_break()

    def save_document(self, filename: str = "content_analysis.docx"):
        """Save document with improved formatting.
        
        Args:
            filename: Name of the output document file
        """
        output_path = self.storage_dir / filename

        # Set page margins
        for section in self.document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        self.document.save(str(output_path))
        logging.info(f"Document saved to {output_path}")


class AsyncWebScraper:
    """Asynchronous web scraper with modular components.
    
    This class orchestrates all scraper components to efficiently crawl
    websites with proper error handling, resource management, and content
    processing.
    
    Attributes:
        config: Scraper configuration
        storage: Backend for content storage
        url_manager: Component for URL validation and tracking
        page_fetcher: Component for fetching web pages
        rate_limiter: Component for rate limiting requests
        circuit_breaker: Component for handling failing domains
        memory_monitor: Component for monitoring memory usage
        proxy_manager: Component for managing proxies
        content_processor: Component for processing content
        content_analyzer: Component for analyzing content
        processed_urls: Counter for processed URLs
        failed_urls: Counter for failed URLs
    """
    
    def __init__(self, config: ScraperConfig, storage: StorageBackend):
        """Initialize scraper with modular components.
        
        Args:
            config: Scraper configuration
            storage: Storage backend for content
        """
        self.config = config
        self.storage = storage
        
        # Initialize components
        self.url_manager = URLManager(config)
        self.page_fetcher = PageFetcher(config)
        self.rate_limiter = RateLimiter(calls=10, period=1.0)
        self.circuit_breaker = CircuitBreaker()
        self.memory_monitor = MemoryMonitor(config.memory_limit_mb)
        self.proxy_manager = ProxyManager(config.proxies)
        self.content_processor = ContentProcessor(
            char_limit=config.max_file_size_mb * 1024 * 1024,
            enable_ai_processing=False  # Disabled by default
        )
        self.content_analyzer = RealTimeContentAnalyzer(config.output_dir)
        
        # Initialize counters
        self.processed_urls = 0
        self.failed_urls = 0
        
        # Setup logging
        self._setup_logging()
        
        # Track visited URLs
        self.visited_urls = set()

    def _setup_logging(self):
        """Setup logging configuration."""
        log_file = self.config.output_dir / "scraper.log"
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(levelname)s - %(message)s",
            handlers=[
                logging.FileHandler(log_file), 
                logging.StreamHandler()
            ],
        )

    async def scrape_page(
        self, url: str, depth: int, session: aiohttp.ClientSession
    ) -> Tuple[Optional[str], List[Tuple[str, str]]]:
        """Scrape a single page using modular components.
        
        Args:
            url: URL to scrape
            depth: Current depth from root URL
            session: aiohttp session for HTTP requests
            
        Returns:
            Tuple of (content, links) where content is the scraped text
            and links is a list of (url, text) tuples for discovered links
        """
        # Check memory usage
        is_within_limit, current_usage = self.memory_monitor.check_memory()
        if not is_within_limit:
            logging.warning(f"Memory usage too high ({current_usage:.1f}MB). Skipping {url}")
            return None, []

        # Validate URL
        if not self.url_manager.is_valid_url(url, depth):
            logging.warning(f"Invalid URL or maximum depth reached: {url}")
            return None, []

        # Check circuit breaker
        domain = urlparse(url).netloc
        if await self.circuit_breaker.is_open(domain):
            logging.warning(f"Circuit breaker open for domain: {domain}")
            return None, []

        # Try to fetch page with retries
        for attempt in range(self.config.max_retries):
            try:
                await self.rate_limiter.acquire()
                headers = {"User-Agent": self.config.user_agent}
                
                # Get proxy if configured
                proxy = await self.proxy_manager.get_next_proxy()
                if proxy:
                    headers["Proxy-Authorization"] = proxy

                # Fetch page content
                content = await self.page_fetcher.fetch_page(url, session, headers)
                if not content:
                    await self.circuit_breaker.record_failure(domain)
                    if attempt < self.config.max_retries - 1:
                        await asyncio.sleep(self.config.delay * (2 ** attempt))
                        continue
                    return None, []

                # Check if content needs processing
                if not self.storage.should_process_url(url, content):
                    logging.info(f"Content unchanged for URL: {url}")
                    return None, []

                # Mark URL as visited and record success
                self.url_manager.mark_visited(url)
                await self.circuit_breaker.record_success(domain)

                # Parse content
                soup = BeautifulSoup(content, "html.parser")
                
                # Extract text content
                text_content = soup.get_text(separator=" ", strip=True)

                # Check content size
                if len(text_content.encode('utf-8')) > self.config.max_file_size_mb * 1024 * 1024:
                    logging.warning(f"Content too large for {url}")
                    return None, []

                # Extract links
                links = [
                    (urljoin(url, a["href"]), a.text.strip())
                    for a in soup.find_all("a", href=True)
                    if self.url_manager.is_valid_url(urljoin(url, a["href"]), depth + 1)
                ]

                # Add sitemap links
                links.extend(self.url_manager.get_sitemap_links(url, depth))

                return text_content, links

            except asyncio.TimeoutError:
                await self.circuit_breaker.record_failure(domain)
                if attempt < self.config.max_retries - 1:
                    await asyncio.sleep(self.config.delay * (2 ** attempt))
                    continue
                logging.error(f"Timeout error for {url} after {attempt + 1} attempts")
                return None, []

            except Exception as e:
                await self.circuit_breaker.record_failure(domain)
                if attempt < self.config.max_retries - 1:
                    await asyncio.sleep(self.config.delay * (2 ** attempt))
                    continue
                logging.error(f"Error scraping {url}: {str(e)}")
                return None, []

        return None, []

    async def process_url(
        self, url: str, depth: int, session: aiohttp.ClientSession
    ) -> Set[str]:
        """Process a URL by scraping, analyzing, and extracting links.
        
        Args:
            url: URL to process
            depth: Current depth from root URL
            session: aiohttp session for HTTP requests
            
        Returns:
            Set of discovered URLs to process next
        """
        try:
            content, links = await self.scrape_page(url, depth, session)
            if not content:
                self.failed_urls += 1
                logging.warning(f"Failed to process URL: {url}")
                return set()

            # Process content using transformer model
            processed_content = self.content_processor.process_text(content)
            
            # Extract and clean sentences
            sentences = self.extract_sentences(processed_content)
            cleaned_content = "\n".join(sentences)
            
            # Add to content buffer
            buffer = self.content_processor.add_text(cleaned_content, url)
            if buffer:
                # Save accumulated content when buffer is full
                self.storage.save_content(f"batch_{self.content_processor.file_counter}", buffer)
            
            # Process content in real-time analyzer
            self.content_analyzer.process_content(url, cleaned_content)
                
            self.processed_urls += 1
            logging.info(
                f"URL successfully processed: {url} "
                f"(total {self.processed_urls}, failed {self.failed_urls})"
            )

            return {link for link, _ in links if link not in self.visited_urls}
            
        except Exception as e:
            self.failed_urls += 1
            logging.error(f"Error processing URL {url}: {str(e)}")
            return set()

    def extract_sentences(self, text: str) -> List[str]:
        """Extract valid sentences from text.
        
        Args:
            text: Input text to process
            
        Returns:
            List of extracted sentences
        """
        sentence_pattern = r"(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s(?=[A-Z])"
        sentences = re.split(sentence_pattern, text)
        return [s.strip() for s in sentences if s.strip() and re.match(r"^[A-Z].*[.!?]$", s.strip())]

    async def run(self, start_url: str) -> None:
        """Run the scraper starting from a URL.
        
        Args:
            start_url: Initial URL to start scraping from
        """
        try:
            self.visited_urls = self.storage.load_urls()
            to_visit = {(start_url, 0)}
            start_time = time.time()

            connector = aiohttp.TCPConnector(
                limit=self.config.max_workers,
                verify_ssl=self.config.verify_ssl
            )
            
            # Initialize page fetcher
            await self.page_fetcher.setup()
            
            async with aiohttp.ClientSession(connector=connector) as session:
                while to_visit:
                    current_batch = set()
                    while to_visit and len(current_batch) < self.config.max_workers:
                        url, depth = to_visit.pop()
                        current_batch.add((url, depth))

                    tasks = [
                        self.process_url(url, depth, session)
                        for url, depth in current_batch
                    ]

                    results = await asyncio.gather(*tasks)

                    for new_urls in results:
                        to_visit.update((url, depth + 1) for url in new_urls)

                    self.storage.save_urls(self.visited_urls)
                    
                    # Save any remaining content in buffer
                    remaining_content = self.content_processor.get_current_buffer()
                    if remaining_content:
                        self.storage.save_content(
                            f"batch_{self.content_processor.file_counter}", 
                            remaining_content
                        )
                    
                    elapsed_time = time.time() - start_time
                    logging.info(
                        f"Progress: {self.processed_urls} processed, {self.failed_urls} failed, "
                        f"Time: {elapsed_time:.2f}s"
                    )
                    
            elapsed_time = time.time() - start_time
            logging.info(
                f"Scraping completed. Total: {self.processed_urls} processed, "
                f"{self.failed_urls} failed, Time: {elapsed_time:.2f}s"
            )
            
            # Save final analysis
            self.content_analyzer.save_document()
            
        except Exception as e:
            logging.error(f"Error during scraping: {str(e)}")
            raise
        finally:
            # Ensure we save any remaining content
            remaining_content = self.content_processor.get_current_buffer()
            if remaining_content:
                self.storage.save_content(
                    f"batch_{self.content_processor.file_counter}", 
                    remaining_content
                )
            
            # Cleanup resources
            await self.page_fetcher.cleanup()


async def main():
    """Main entry point for the scraper."""
    # Load configuration from file or user input
    import argparse
    
    parser = argparse.ArgumentParser(description="Advanced web scraper")
    parser.add_argument("--config", type=str, help="Path to configuration file")
    parser.add_argument("--url", type=str, help="Root URL to scrape")
    parser.add_argument("--output", type=str, default="output", help="Output directory")
    args = parser.parse_args()
    
    if args.config:
        try:
            config = ScraperConfig.from_json_file(args.config)
            logging.info(f"Configuration loaded from {args.config}")
        except Exception as e:
            logging.error(f"Error loading configuration: {e}")
            return
    else:
        root_url = args.url or input("Enter root URL: ").strip()
        config = ScraperConfig(
            root_url=root_url,
            output_dir=Path(args.output),
            max_retries=5,
            delay=3.0,
            max_workers=5,
            max_depth=5,
            enable_javascript=PLAYWRIGHT_AVAILABLE,
            memory_limit_mb=1000,
            max_file_size_mb=50,
            proxies=[],
        )

    try:
        # Create output directory
        config.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Save configuration for reference
        config_path = config.output_dir / "config.json"
        config.save_to_json(config_path)
        logging.info(f"Configuration saved to {config_path}")
        
        # Initialize storage and scraper
        storage = FileStorage(config.output_dir)
        scraper = AsyncWebScraper(config, storage)

        logging.info(f"Starting scraper with configuration: {config.to_dict()}")
        await scraper.run(config.root_url)
        logging.info("Scraping completed")
        
        # Save final analysis document
        scraper.content_analyzer.save_document()
        logging.info("Content analysis document saved")
        
    except Exception as e:
        logging.error(f"Fatal error: {e}")
    

if __name__ == "__main__":
    asyncio.run(main())
